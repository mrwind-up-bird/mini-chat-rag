"""Tests for app.services.extract — text extraction from uploaded files."""

import pytest

from app.services.extract import extract_text


def test_extract_txt():
    text = extract_text("readme.txt", b"Hello, world!")
    assert text == "Hello, world!"


def test_extract_md():
    text = extract_text("notes.md", b"# Heading\n\nParagraph")
    assert text == "# Heading\n\nParagraph"


def test_extract_csv():
    csv_bytes = b"name,age\nAlice,30\nBob,25"
    text = extract_text("data.csv", csv_bytes)
    assert "Alice" in text
    assert "Bob" in text


def test_extract_pdf():
    """Create a minimal valid PDF and extract text from it."""
    from io import BytesIO

    from pypdf import PdfWriter

    writer = PdfWriter()
    page = writer.add_blank_page(width=72, height=72)

    # Add text annotation via reportlab-free approach
    from pypdf.generic import (
        DictionaryObject,
        NameObject,
    )

    # Build a minimal content stream with text
    content = b"BT /F1 12 Tf 10 50 Td (Hello PDF) Tj ET"

    # Create a simple font dict
    font_dict = DictionaryObject()
    font_dict.update(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )

    resources = DictionaryObject()
    font_resources = DictionaryObject()
    font_resources[NameObject("/F1")] = font_dict
    resources[NameObject("/Font")] = font_resources

    from pypdf.generic import DecodedStreamObject

    stream = DecodedStreamObject()
    stream.set_data(content)

    page[NameObject("/Resources")] = resources
    page[NameObject("/Contents")] = stream

    buf = BytesIO()
    writer.write(buf)
    pdf_bytes = buf.getvalue()

    text = extract_text("doc.pdf", pdf_bytes)
    assert "Hello PDF" in text


def test_extract_docx():
    """Create a minimal .docx and extract text from it."""
    from io import BytesIO

    from docx import Document

    doc = Document()
    doc.add_paragraph("Hello DOCX")
    doc.add_paragraph("Second paragraph")

    buf = BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    text = extract_text("report.docx", docx_bytes)
    assert "Hello DOCX" in text
    assert "Second paragraph" in text


def test_unsupported_extension():
    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_text("malware.exe", b"\x00\x01\x02")
